"""Document text extraction for PDF, Word, Markdown, and plain text files.

Each extractor returns a list of "blocks" (paragraphs). Blocks are the unit of
translation: keeping paragraphs separate gives the small MT model bounded,
coherent inputs and lets us show progress per block.
"""

from __future__ import annotations

import os
import re
from typing import List


DOC_EXTENSIONS = {".pdf", ".docx", ".doc", ".md", ".markdown", ".txt"}
IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".bmp", ".webp", ".tif", ".tiff", ".gif"}
SUPPORTED_EXTENSIONS = DOC_EXTENSIONS | IMAGE_EXTENSIONS


def is_image(path: str) -> bool:
    return os.path.splitext(path)[1].lower() in IMAGE_EXTENSIONS


def extract_blocks(path: str) -> List[str]:
    """Return a list of text blocks (paragraphs) from a supported document.

    Images are not handled here — they require OCR via the model backend.
    """
    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf":
        return _extract_pdf(path)
    if ext in (".docx", ".doc"):
        return _extract_docx(path)
    if ext in (".md", ".markdown"):
        return _extract_text(path)
    if ext == ".txt":
        return _extract_text(path)
    if ext in IMAGE_EXTENSIONS:
        raise ValueError("图片需通过 OCR 处理，请使用 OCR 流程")
    raise ValueError(f"不支持的文件类型: {ext}")


def split_text_into_blocks(raw: str) -> List[str]:
    """Split free-form text (e.g. OCR output) into paragraph blocks."""
    parts = re.split(r"\n\s*\n", raw.replace("\r\n", "\n"))
    return _clean_blocks(parts)


def _clean_blocks(blocks: List[str]) -> List[str]:
    """Normalise whitespace and drop empty blocks."""
    out: List[str] = []
    for b in blocks:
        b = b.replace("\r\n", "\n").strip()
        if b:
            out.append(b)
    return out


def _extract_pdf(path: str) -> List[str]:
    try:
        import pymupdf as fitz  # PyMuPDF >= 1.24 canonical name
    except ImportError:
        import fitz  # older alias

    blocks: List[str] = []
    with fitz.open(path) as doc:
        for page in doc:
            # "blocks" gives (x0, y0, x1, y1, text, block_no, block_type)
            page_blocks = page.get_text("blocks")
            # sort top-to-bottom, then left-to-right
            page_blocks.sort(key=lambda b: (round(b[1], 1), b[0]))
            for pb in page_blocks:
                text = pb[4]
                if not text or not text.strip():
                    continue
                # join intra-paragraph line breaks but keep the paragraph whole
                merged = _dehyphenate(text)
                blocks.append(merged)
    return _clean_blocks(blocks)


def _dehyphenate(text: str) -> str:
    """Collapse soft line wraps inside a PDF block into spaces."""
    text = text.replace("\r\n", "\n")
    # remove hyphenation at line ends: "exam-\nple" -> "example"
    text = re.sub(r"-\n(\w)", r"\1", text)
    # turn remaining single newlines into spaces (paragraph stays one block)
    text = re.sub(r"\n+", " ", text)
    return re.sub(r"[ \t]+", " ", text).strip()


def _extract_docx(path: str) -> List[str]:
    import docx

    document = docx.Document(path)
    blocks: List[str] = []
    for para in document.paragraphs:
        if para.text and para.text.strip():
            blocks.append(para.text)
    # also pull text out of tables
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text and cell.text.strip():
                    blocks.append(cell.text)
    return _clean_blocks(blocks)


def _extract_text(path: str) -> List[str]:
    raw = _read_text_file(path)
    return split_text_into_blocks(raw)


def _read_text_file(path: str) -> str:
    for enc in ("utf-8-sig", "utf-8", "gb18030", "latin-1"):
        try:
            with open(path, "r", encoding=enc) as f:
                return f.read()
        except (UnicodeDecodeError, UnicodeError):
            continue
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        return f.read()
